from docx import Document
import re
import os
import requests


def extract_docx_paths(folder_path):
    docx_paths = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.docx'):
                full_path = os.path.join(root, file)
                docx_paths.append(full_path)
    filtered_paths = [path for path in docx_paths if not os.path.basename(path).startswith('~$')]
    return filtered_paths

def extract_all_text(docx_path):
    """
    Extracts all text from a .docx file, including from tables, into a formatted and cleaned string.
    Preserves paragraph structure with double newlines and cleans up excess whitespace.
    
    Args:
        docx_path (str): Path to the .docx file.
    
    Returns:
        str: The cleaned and formatted text.
    """
    doc = Document(docx_path)
    full_text = []
    
    # Helper function to extract text from paragraphs (used for body and table cells)
    def extract_from_paragraphs(paragraphs):
        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
    
    # Extract from main document body
    extract_from_paragraphs(doc.paragraphs)
    
    # Extract from tables (appended after body; for better order, consider doc structure if needed)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extract_from_paragraphs(cell.paragraphs)
    
    # Join paragraphs with double newlines for formatting
    formatted_text = '\n\n'.join(full_text)
    
    # Clean up: remove multiple consecutive newlines and excess whitespace
    formatted_text = re.sub(r'\n{3,}', '\n\n', formatted_text)  # Replace 3+ newlines with 2
    formatted_text = re.sub(r'\s{2,}', ' ', formatted_text)    # Replace multiple spaces with single space
    
    return formatted_text.strip()


def extract_highlighted_text(docx_path):
    """
    Extracts all highlighted text from a .docx file.
    
    Args:
        docx_path (str): Path to the .docx file.
    
    Returns:
        list: A list of strings containing highlighted text segments.
    """
    doc = Document(docx_path)
    highlighted_texts = []
    
    # Helper function to extract from paragraphs (used for main body and tables)
    def extract_from_paragraphs(paragraphs):
        for paragraph in paragraphs:
            for run in paragraph.runs:
                if run.font.highlight_color is not None:
                    highlighted_texts.append(run.text.strip())
    
    # Extract from main document body
    extract_from_paragraphs(doc.paragraphs)
    
    # Extract from tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extract_from_paragraphs(cell.paragraphs)
    
    return highlighted_texts

def call_grok(api_key: str, system_prompt: str, user_prompt: str, model: str = "grok-3-fast") -> dict:
    """
    Calls the xAI Grok API chat completions endpoint with the specified payload.
    """
    url = "https://api.x.ai/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    payload = {
        "messages": [
            {
                "role": "system",
                "content": f"{system_prompt}"
            },
            {
                "role": "user",
                "content": f"{user_prompt}"
            }
        ],
        "model": model,
        "stream": False,
        "temperature": 0
    }
    response = requests.post(url, headers=headers, json=payload)
    response.raise_for_status()  # Raise an exception for HTTP errors
    return response.json()